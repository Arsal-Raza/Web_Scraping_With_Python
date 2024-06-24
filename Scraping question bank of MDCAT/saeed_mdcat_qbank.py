from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from bs4 import BeautifulSoup
import requests
from docx import Document

# Initialize the Chrome driver
chrome_driver_path = r"D:\chrome driver\chromedriver-win64\chromedriver.exe"
service = Service(chrome_driver_path)
options = Options()
driver = webdriver.Chrome(service=service, options=options)

try:
    # Open the course page
    driver.get("https://saeedmdcatlms.com/courses/alpha-2-0-cts-session/")
    driver.maximize_window()

    # Click the "Enroll now" button
    enroll_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Enroll now')]"))
    )
    enroll_button.click()
    time.sleep(3)

    # Fill the login form and submit
    username = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.NAME, "log"))
    )
    password = driver.find_element(By.NAME, "pwd")
    login_button = driver.find_element(By.XPATH, "//button[contains(text(), 'Sign In')]")

    username.send_keys("username")
    password.send_keys("password")
    login_button.click()
    print('Login successful.')

    # Click the "Start Learning" button
    start_learning_button = WebDriverWait(driver, 10).until(
        EC.visibility_of_element_located((By.XPATH, "//a[contains(text(), 'Start Learning')]"))
    )
    driver.execute_script("arguments[0].scrollIntoView(true);", start_learning_button)
    time.sleep(1)
    start_learning_button.click()
    print('Course started...')
    print()
    time.sleep(3)

    # Get all the paper elements and their URLs
    papers = driver.find_elements(By.CLASS_NAME, "tutor-course-topic-item-lesson")
    paper_urls = [paper.find_element(By.TAG_NAME, "a").get_attribute("href") for paper in papers]

    # Initialize a document to save the content
    doc = Document()

    # Loop through each paper URL
    for paper_url in paper_urls:
        try:
            # Open the paper URL
            driver.get(paper_url)
            time.sleep(5)

            # Get the title of the paper
            paper_title = driver.find_element(By.TAG_NAME, "h2").text.strip()

            # Click the "Quiz-summary" button
            quiz_summary_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.NAME, "quizSummary"))
            )
            quiz_summary_button.click()
            print(f'Clicked on quiz summary button for "{paper_title}"')
            time.sleep(3)

            # Click the "Finish quiz" button
            finish_quiz_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.NAME, "endQuizSummary"))
            )
            finish_quiz_button.click()
            print('Clicked on finish quiz button')
            time.sleep(3)

            # Click the "View questions" button
            view_questions_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.NAME, "reShowQuestion"))
            )
            view_questions_button.click()
            print('Clicked on view questions button')
            time.sleep(5)

            # Scraping questions and answers
            questions = driver.find_elements(By.CLASS_NAME, "wpProQuiz_listItem")
            print(f'Got {len(questions)} questions for "{paper_title}"...')
            
            # Add paper title to the document
            doc.add_heading(paper_title, level=1)
            
            # Add each question and answer to the document
            for question in questions:
                question_no = question.find_element(By.CSS_SELECTOR, "h5.wpProQuiz_header span").text.strip()
                question_text = question.find_element(By.CLASS_NAME, "wpProQuiz_question_text").text.strip()
                options = question.find_elements(By.CLASS_NAME, "wpProQuiz_questionListItem")

                doc.add_paragraph(f"{question_no}. {question_text}", style='List Number')
                for option in options:
                    option_text = option.find_element(By.TAG_NAME, "label").text.strip()
                    is_correct = "wpProQuiz_answerCorrect" in option.get_attribute("class")
                    doc.add_paragraph(f" - {option_text}", style='Normal')
                    if is_correct:
                        correct_answer = option_text
                doc.add_paragraph(f"\nCorrect Answer: {correct_answer}\n" if 'correct_answer' in locals() else "\nCorrect Answer not found.\n", style='Normal')

            doc.add_paragraph("")  # Add empty paragraph for spacing between papers

            print("Paper processing completed.\n")

        except Exception as e:
            print(f"An error occurred while processing paper: {str(e)}")
            # Continue to the next paper even if there's an error

except Exception as e:
    print(f"An error occurred during script execution: {str(e)}")

finally:
    # Save the document
    doc.save('papers_questions_answers.docx')

    # Close the driver
    driver.quit()
